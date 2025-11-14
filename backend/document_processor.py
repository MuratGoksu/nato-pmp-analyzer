"""
Document Processing Module
"""

import PyPDF2
from docx import Document
from pathlib import Path
from typing import Dict, List, Optional
import re
from datetime import datetime


class DocumentProcessor:
    """Process PDF and DOCX documents"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def process_document(self, file_path: str, file_name: str) -> Dict:
        file_extension = Path(file_name).suffix.lower()
        
        result = {
            'file_name': file_name,
            'file_type': file_extension,
            'processed_at': datetime.now().isoformat(),
            'success': False,
            'error': None,
            'text': '',
            'metadata': {}
        }
        
        try:
            if file_extension == '.pdf':
                result['text'] = self._extract_pdf_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                result['text'] = self._extract_docx_text(file_path)
            else:
                result['error'] = f"Unsupported format: {file_extension}"
                return result
            
            result['metadata'] = self._extract_metadata(result['text'], file_name)
            result['success'] = True
            
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def _extract_pdf_text(self, file_path: str) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n\n"
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        text = ""
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        return text.strip()
    
    def _extract_metadata(self, text: str, file_name: str) -> Dict:
        return {
            'project_name': self._extract_project_name(text, file_name),
            'budget': self._extract_budget(text),
            'dates': self._extract_dates(text),
            'stakeholders': self._extract_stakeholders(text),
            'status': self._extract_status(text),
            'word_count': len(text.split()),
            'char_count': len(text)
        }
    
    def _extract_project_name(self, text: str, file_name: str) -> Optional[str]:
        patterns = [
            r'Project\s+Name\s*:\s*([^\n]+)',
            r'Project\s+Title\s*:\s*([^\n]+)',
            r'Project\s*:\s*([^\n]+)',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return Path(file_name).stem
    
    def _extract_budget(self, text: str) -> List[str]:
        budgets = []
        patterns = [
            r'€\s*[\d,\.]+\s*[MmKk]?',
            r'[\d,\.]+\s*(?:million|Million|M€|EUR)',
            r'Budget\s*:\s*€?\s*[\d,\.]+',
        ]
        for pattern in patterns:
            budgets.extend(re.findall(pattern, text))
        return list(set(budgets))
    
    def _extract_dates(self, text: str) -> List[str]:
        dates = []
        patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',
            r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}',
        ]
        for pattern in patterns:
            dates.extend(re.findall(pattern, text, re.IGNORECASE))
        return list(set(dates))[:10]
    
    def _extract_stakeholders(self, text: str) -> List[str]:
        stakeholders = []
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        stakeholders.extend(re.findall(email_pattern, text))
        return list(set(stakeholders))[:10]
    
    def _extract_status(self, text: str) -> Optional[str]:
        text_lower = text.lower()
        if re.search(r'\b(?:status\s*:\s*)?red\b', text_lower):
            return "RED"
        elif re.search(r'\b(?:status\s*:\s*)?amber\b', text_lower):
            return "AMBER"
        elif re.search(r'\b(?:status\s*:\s*)?green\b', text_lower):
            return "GREEN"
        
        risk_keywords = ['critical', 'high risk', 'delayed', 'overbudget']
        if any(keyword in text_lower for keyword in risk_keywords):
            return "RED"
        
        return "GREEN"
    
    def get_text_preview(self, text: str, max_length: int = 500) -> str:
        if len(text) <= max_length:
            return text
        return text[:max_length] + "..."
